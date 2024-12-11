import torch
import torch.nn as nn
import torch.nn.functional as F
from torch.utils.data import DataLoader, Dataset, random_split
from docx import Document  # To read Word documents
import dash
from dash import html, dcc, Input, Output, State
import dash_bootstrap_components as dbc
from transformers import pipeline
import random
import os
import base64
from io import BytesIO

class LSTMModel(nn.Module):
    def __init__(self, vocab_size, embed_size, hidden_size, num_layers):
        super(LSTMModel, self).__init__()
        self.embedding = nn.Embedding(vocab_size, embed_size)
        self.lstm = nn.LSTM(embed_size, hidden_size, num_layers, batch_first=True)
        self.fc = nn.Linear(hidden_size, vocab_size)

    def forward(self, x, hidden):
        embed = self.embedding(x)
        out, hidden = self.lstm(embed, hidden)
        out = self.fc(out[:, -1, :])  # Take the last output
        return out, hidden

def generate_sentence_lstm(model, start_words, vocab, reverse_vocab, max_length=10, temperature=1.0, top_k=5):
    model.eval()
    hidden = None
    words = start_words.split()
    sentence = words[:]

    input_ids = torch.tensor([[vocab.get(word, vocab['<unk>']) for word in words]], dtype=torch.long)
    for _ in range(max_length - len(words)):
        with torch.no_grad():
            output, hidden = model(input_ids, hidden)
            logits = output.squeeze() / temperature

            # Top-k sampling
            top_k_logits, top_k_indices = torch.topk(logits, k=top_k)
            probabilities = F.softmax(top_k_logits, dim=-1).numpy()
            predicted_index = random.choices(top_k_indices.numpy(), probabilities)[0]

            next_word = reverse_vocab.get(predicted_index, '<unk>')
            if next_word == '<eos>' or next_word in sentence:
                break
            sentence.append(next_word)
            input_ids = torch.tensor([[predicted_index]], dtype=torch.long)
    return " ".join(sentence)

def read_word_document(file_path):
    document = Document(file_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return " ".join(text)

def train_model_on_new_data(file_content, model, optimizer, criterion, vocab, reverse_vocab):
    # Tokenize new content
    words = file_content.split()
    tokenized_data = [vocab.get(word, vocab['<unk>']) for word in words]

    # Prepare input-output pairs with a sliding window
    inputs, targets = [], []
    for i in range(len(tokenized_data) - 1):
        inputs.append(tokenized_data[i])  # Current word
        targets.append(tokenized_data[i + 1])  # Next word

    inputs = torch.tensor(inputs, dtype=torch.long).view(-1, 1)  # Shape: (sequence_length, 1)
    targets = torch.tensor(targets, dtype=torch.long)  # Shape: (sequence_length,)

    model.train()
    hidden = None
    optimizer.zero_grad()
    output, hidden = model(inputs, hidden)  # Forward pass
    loss = criterion(output, targets)  # Compute loss
    loss.backward()
    optimizer.step()
    return loss.item()

# Sentiment Analysis Pipeline
sentiment_analyzer = pipeline("sentiment-analysis")

def analyze_sentiment(input_text):
    result = sentiment_analyzer(input_text)[0]
    return result['label'], result['score']

# Load data and initialize model
file_path = "data.docx"
try:
    text = read_word_document(file_path)
except Exception as e:
    print(f"Error reading Word document: {e}")
    exit(1)

# Vocabulary setup
unique_words = set(text.split())
vocab = {word: idx for idx, word in enumerate(unique_words, start=1)}
vocab['<unk>'] = 0  # For unknown words
vocab['<eos>'] = len(vocab)
reverse_vocab = {idx: word for word, idx in vocab.items()}

# Model setup
embed_size = 128
hidden_size = 256
num_layers = 2
model = LSTMModel(len(vocab), embed_size, hidden_size, num_layers)
optimizer = torch.optim.Adam(model.parameters(), lr=0.001)
criterion = nn.CrossEntropyLoss()

# Dash app
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])
app.layout = dbc.Container([
    html.Div(
        [
            html.H1("Human-like Sentence Generator", className="text-center my-4", style={"color": "#28A745", "fontWeight": "bold"}),
            dbc.Card([
                dbc.CardBody([
                    dbc.Row([
                        dbc.Col([
                            dcc.Input(id="input-text", type="text", placeholder="Enter starting words...", className="form-control mb-3"),
                            dbc.Button("Generate Sentence", id="predict-button", color="primary", className="mb-3 w-100"),
                            html.Div(id="prediction-output", className="alert alert-info", style={"display": "none"}),
                            html.Div(id="sentiment-output", className="alert alert-warning", style={"display": "none"}),
                        ], width=6),
                    ]),
                    dbc.Row([
                        dbc.Col([
                            dcc.Upload(
                                id="upload-data",
                                children=html.Div([
                                    "Drag and Drop or ",
                                    html.A("Select a File")
                                ]),
                                style={
                                    "width": "100%",
                                    "height": "60px",
                                    "lineHeight": "60px",
                                    "borderWidth": "1px",
                                    "borderStyle": "dashed",
                                    "borderRadius": "5px",
                                    "textAlign": "center",
                                    "margin": "10px 0",
                                },
                                multiple=False
                            ),
                            html.Div(id="upload-output", className="alert alert-success", style={"display": "none"})
                        ], width=12),
                    ])
                ])
            ], style={"boxShadow": "0 4px 12px rgba(0, 0, 0, 0.1)", "borderRadius": "10px", "backgroundColor": "#1A1A1D", "color": "white"}),
            dbc.Card([
                dbc.CardBody([
                    html.H4("Video Preview", className="card-title", style={"color": "white"}),
                    html.Iframe(
                        src="https://www.youtube.com/embed/Ejw2eZhvHHY",
                        style={"width": "100%", "height": "315px", "border": "none"}
                    )
                ])
            ], style={"boxShadow": "0 4px 12px rgba(0, 0, 0, 0.1)", "borderRadius": "10px", "backgroundColor": "#1A1A1D", "color": "white", "marginTop": "20px"})
        ]
    )
], style={"backgroundColor": "black", "height": "100vh", "paddingTop": "50px"})

@app.callback(
    [Output("prediction-output", "children"), Output("prediction-output", "style"),
     Output("sentiment-output", "children"), Output("sentiment-output", "style")],
    Input("predict-button", "n_clicks"),
    State("input-text", "value")
)
def handle_sentence_generation(n_clicks, input_text):
    if n_clicks is None or not input_text:
        return "", {"display": "none"}, "", {"display": "none"}

    try:
        sentiment_label, sentiment_score = analyze_sentiment(input_text)
        sentiment_text = f"Sentiment: {sentiment_label} (Confidence: {sentiment_score:.2f})"

        generated_sentence = generate_sentence_lstm(model, input_text, vocab, reverse_vocab, temperature=0.7, top_k=5)
        return (f"Generated Sentence: {generated_sentence}", {"display": "block"},
                sentiment_text, {"display": "block"})

    except Exception as e:
        return (f"Error: {str(e)}", {"display": "block", "color": "red"}, "", {"display": "none"})

@app.callback(
    Output("upload-output", "children"),
    Output("upload-output", "style"),
    Input("upload-data", "contents"),
    State("upload-data", "filename")
)
def handle_file_upload(contents, filename):
    if contents is None:
        return "", {"display": "none"}

    try:
        # Decode file contents
        content_type, content_string = contents.split(",")
        decoded = base64.b64decode(content_string)

        # Load the .docx content
        document = Document(BytesIO(decoded))
        file_content = " ".join([para.text for para in document.paragraphs])

        # Train model on new data
        loss = train_model_on_new_data(file_content, model, optimizer, criterion, vocab, reverse_vocab)
        return f"File '{filename}' uploaded and model trained. Loss: {loss:.4f}", {"display": "block"}

    except Exception as e:
        return f"Error processing file: {str(e)}", {"display": "block", "color": "red"}

if __name__ == "__main__":
    app.run_server(debug=True)
